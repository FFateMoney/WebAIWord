import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .paragraph_parser import parse_paragraph_block
from .style_parser import parse_styles
from .table_parser import parse_table_block


def _parse_meta(doc) -> dict:
    section = doc.sections[0]
    return {
        "page": {
            "size": "custom",
            "width": section.page_width.twips,
            "height": section.page_height.twips,
            "orientation": "landscape" if section.page_width > section.page_height else "portrait",
            "margin": {
                "top": section.top_margin.twips,
                "bottom": section.bottom_margin.twips,
                "left": section.left_margin.twips,
                "right": section.right_margin.twips,
            },
        },
        "default_style": "Normal",
        "language": "zh-CN",
    }


def _is_toc_sdt(sdt_el) -> bool:
    _tag_sdtPr = qn("w:sdtPr")
    _tag_docPartObj = qn("w:docPartObj")
    _tag_gallery = qn("w:docPartGallery")

    sdtPr = sdt_el.find(_tag_sdtPr)
    if sdtPr is not None:
        docPartObj = sdtPr.find(_tag_docPartObj)
        if docPartObj is not None:
            gallery = docPartObj.find(_tag_gallery)
            if gallery is not None:
                val = gallery.get(qn("w:val"), "")
                if "Table of Contents" in val:
                    return True

    sdt_content = sdt_el.find(qn("w:sdtContent"))
    if sdt_content is not None:
        for instrText in sdt_content.iter(qn("w:instrText")):
            if instrText.text and instrText.text.strip().upper().startswith("TOC"):
                return True
    return False


def _parse_toc_block(sdt_el, doc, block_id) -> dict:
    sdt_content = sdt_el.find(qn("w:sdtContent"))
    _tag_p = qn("w:p")
    _tag_fldChar = qn("w:fldChar")
    _tag_instrText = qn("w:instrText")

    instruction_parts: list[str] = []
    in_field = False
    for el in (sdt_content if sdt_content is not None else []):
        if el.tag != _tag_p:
            continue
        for r_el in el.iter(qn("w:r")):
            fc = r_el.find(_tag_fldChar)
            if fc is not None:
                ft = fc.get(qn("w:fldCharType"))
                if ft == "begin":
                    in_field = True
                    continue
                if ft in ("separate", "end"):
                    in_field = False
                    continue
            if in_field:
                it = r_el.find(_tag_instrText)
                if it is not None and it.text:
                    instruction_parts.append(it.text)

    instruction = "".join(instruction_parts).strip()
    if not instruction:
        instruction = 'TOC \\o "1-3" \\h \\z \\u'

    title = None
    if sdt_content is not None:
        for child in sdt_content:
            if child.tag != _tag_p:
                continue
            has_fld_begin = any(
                fc.get(qn("w:fldCharType")) == "begin"
                for fc in child.iter(_tag_fldChar)
            )
            if has_fld_begin:
                break
            paragraph = Paragraph(child, doc)
            if paragraph.text.strip():
                title = parse_paragraph_block(paragraph, f"{block_id}.title")
                break

    block: dict = {
        "id": block_id,
        "type": "TOC",
        "instruction": instruction,
    }
    if title:
        block["title"] = title
    return block


def parse_docx(input_path: str | Path, output_dir: str | Path | None = None) -> dict:
    doc = Document(str(input_path))
    body = []
    p_i = 0
    t_i = 0
    toc_i = 0

    _tag_sdt = qn("w:sdt")
    _tag_sdt_content = qn("w:sdtContent")

    def _process_body_element(child):
        nonlocal p_i, t_i
        tag = child.tag.split("}")[-1]
        if tag == "p":
            paragraph = next((p for p in doc.paragraphs if p._p is child), None)
            if paragraph is None:
                paragraph = Paragraph(child, doc)
            body.append(parse_paragraph_block(paragraph, f"p{p_i}"))
            p_i += 1
        elif tag == "tbl":
            table = next((t for t in doc.tables if t._tbl is child), None)
            if table is None:
                table = Table(child, doc)
            body.append(parse_table_block(table, f"t{t_i}"))
            t_i += 1

    for child in doc.element.body:
        if child.tag == _tag_sdt:
            if _is_toc_sdt(child):
                body.append(_parse_toc_block(child, doc, f"toc{toc_i}"))
                toc_i += 1
            else:
                sdt_content = child.find(_tag_sdt_content)
                if sdt_content is not None:
                    for inner in sdt_content:
                        _process_body_element(inner)
        elif child.tag == qn("w:sectPr"):
            continue
        else:
            _process_body_element(child)

    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": _parse_meta(doc),
            "styles": parse_styles(doc),
            "body": body,
            "passthrough": {},
        },
    }

    if output_dir:
        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        (out_dir / "document.ast.json").write_text(json.dumps(ast, ensure_ascii=False, indent=2), encoding="utf-8")
        (out_dir / "media").mkdir(exist_ok=True)

    return ast
