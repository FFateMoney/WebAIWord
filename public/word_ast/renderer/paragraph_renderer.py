import base64
import copy
import io

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.shared import RGBColor, Pt, Twips

from word_ast.utils.units import half_points_to_pt

_ALIGN_FROM_STR = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_raw_rPr(run, raw_rPr: str) -> None:
    try:
        new_rPr = parse_xml(raw_rPr)
    except Exception:
        return
    for tag in (qn("w:rPrChange"),):
        el = new_rPr.find(tag)
        if el is not None:
            new_rPr.remove(el)
    r_el = run._element
    old_rPr = r_el.find(qn("w:rPr"))
    if old_rPr is not None:
        r_el.remove(old_rPr)
    r_el.insert(0, new_rPr)


def _apply_raw_pPr(paragraph, raw_pPr: str) -> None:
    try:
        new_pPr = parse_xml(raw_pPr)
    except Exception:
        return
    for tag in (qn("w:numPr"), qn("w:sectPr"), qn("w:pPrChange")):
        el = new_pPr.find(tag)
        if el is not None:
            new_pPr.remove(el)

    if new_pPr.find(qn("w:pStyle")) is None:
        current_pPr = paragraph._element.find(qn("w:pPr"))
        if current_pPr is not None:
            existing_pStyle = current_pPr.find(qn("w:pStyle"))
            if existing_pStyle is not None:
                new_pPr.insert(0, copy.deepcopy(existing_pStyle))

    p_el = paragraph._element
    old_pPr = p_el.find(qn("w:pPr"))
    if old_pPr is not None:
        p_el.remove(old_pPr)
    p_el.insert(0, new_pPr)


def _apply_paragraph_style(paragraph, style_id: str | None, styles: dict | None):
    if not style_id:
        return

    candidates = []
    if isinstance(styles, dict):
        style_def = styles.get(style_id)
        style_name = style_def.get("name") if isinstance(style_def, dict) else None
        if style_name:
            candidates.append(style_name)
    candidates.append(style_id)

    for candidate in candidates:
        try:
            paragraph.style = candidate
            return
        except Exception:
            continue


def _apply_paragraph_format(paragraph, fmt: dict):
    if not fmt:
        return
    if "_raw_pPr" in fmt:
        _apply_raw_pPr(paragraph, fmt["_raw_pPr"])
        return
    pf = paragraph.paragraph_format
    alignment = fmt.get("alignment")
    if alignment and alignment in _ALIGN_FROM_STR:
        pf.alignment = _ALIGN_FROM_STR[alignment]
    if "indent_left" in fmt:
        pf.left_indent = Twips(fmt["indent_left"])
    if "indent_right" in fmt:
        pf.right_indent = Twips(fmt["indent_right"])
    if "indent_first_line" in fmt:
        pf.first_line_indent = Twips(fmt["indent_first_line"])
    if "space_before" in fmt:
        pf.space_before = Twips(fmt["space_before"])
    if "space_after" in fmt:
        pf.space_after = Twips(fmt["space_after"])


def _apply_run_overrides(run, overrides: dict) -> None:
    if "_raw_rPr" in overrides:
        _apply_raw_rPr(run, overrides["_raw_rPr"])
        return
    if "bold" in overrides:
        run.bold = overrides["bold"]
    if "italic" in overrides:
        run.italic = overrides["italic"]
    if "underline" in overrides:
        run.underline = overrides["underline"]
    if "size" in overrides:
        size_pt = half_points_to_pt(overrides["size"])
        if size_pt is not None:
            run.font.size = Pt(size_pt)
    if "color" in overrides and overrides["color"].startswith("#"):
        hex_color = overrides["color"][1:]
        if len(hex_color) == 6:
            run.font.color.rgb = RGBColor.from_string(hex_color)
    if overrides.get("font_ascii"):
        run.font.name = overrides["font_ascii"]
    if overrides.get("font_east_asia"):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), overrides["font_east_asia"])


def render_paragraph(doc, block: dict, styles: dict | None = None):
    paragraph = doc.add_paragraph()
    _apply_paragraph_style(paragraph, block.get("style"), styles)
    _apply_paragraph_format(paragraph, block.get("paragraph_format", {}))

    for piece in block.get("content", []):
        if piece.get("type") == "InlineImage":
            try:
                image_bytes = base64.b64decode(piece["data"])
                run = paragraph.add_run()
                width = piece.get("width")
                height = piece.get("height")
                run.add_picture(
                    io.BytesIO(image_bytes),
                    width=Twips(width) if width else None,
                    height=Twips(height) if height else None,
                )
            except (KeyError, ValueError, OSError):
                pass
            continue
        if piece.get("type") != "Text":
            continue
        run = paragraph.add_run(piece.get("text", ""))
        run_overrides = piece.get("overrides", {})
        _apply_run_overrides(run, run_overrides)
